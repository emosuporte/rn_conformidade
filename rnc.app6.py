import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import base64
import os

# Título
st.title("Registro de Não Conformidades")
st.header("POP.ENF.LAB-PC 010")
# Título
#doc.add_heading("Registro de Não Conformidades", level=1)
#doc.add_heading("POP.ENF.LAB-PC 010", level=2)
#st.set_page_config(page_title="Registro de Não Conformidades")

# Dados iniciais
registros = []
df = pd.read_excel('registros_nao_conformidades.xlsx') if 'registros_nao_conformidades.xlsx' in os.listdir() else None
contador_registro = len(df) + 1 if df is not None else 1

# Carregar o modelo de documento (template)
template_path = 'template.docx'

# Função para substituir texto
def docx_replace(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace(cell, old_text, new_text)

# Inicializar variáveis
nao_conformidade_aberta_por = ""
numero_pedido_cliente = ""
tipo_nao_conformidade = ""
descreva_o_fato = ""
acao_corretiva_imediata = ""
responsavel_acao_corretiva = ""

# Formulário de Registro
with st.form(key='registro_form'):
    # Número de Registro
    contador_registro_text = f"{contador_registro}"

    # Data do Registro
    data_registro = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    # Não Conformidade Aberta por
    nao_conformidade_aberta_por = st.text_input("Não Conformidade Aberta por", value=nao_conformidade_aberta_por)

    # Nº Pedido do Cliente
    numero_pedido_cliente = st.text_input("Nº Pedido do Cliente", value=numero_pedido_cliente)

    # Escolha do Tipo de Não Conformidade
    tipo_nao_conformidade = st.selectbox(
        "Tipo de Não Conformidade",
        [
            "Coleta: Troca de paciente",
            "Coleta: Troca de etiquetas",
            "Coleta: Coleta em tubo inadequado",
            "Coleta: Material sem identificação do paciente",
            "Coleta: Material não coleta",
            "Secretaria: Erro de cadastro",
            "Secretaria: Troca de etiquetas nos tubos/frascos do mesmo paciente",
            "Secretaria: Troca de etiquetas nos tubos/frascos de pacientes diferentes",
            "Outro",
            "Área Técnica: Exames não realizados",
            "Área Técnica: Erro na liberação do exame",
            "Área Técnica: Controle interno fora das especificações",
            "Área Técnica: Equipamentos",
            "Área Técnica: Outro"
        ],
        index=0
    )

    # Descreva o Fato
    descreva_o_fato = st.text_area("Descreva o Fato", value=descreva_o_fato)

    # Ação Corretiva Imediata
    acao_corretiva_imediata = st.text_area("Ação Corretiva Imediata", value=acao_corretiva_imediata)

    # Responsável pela Ação Corretiva Imediata
    responsavel_acao_corretiva = st.text_input("Responsável pela Ação Corretiva Imediata", value=responsavel_acao_corretiva)

    # Botão para Registrar Não Conformidade
    submit_button = st.form_submit_button(label='Registrar Não Conformidade')

# Manipulação dos dados
if submit_button:
    novo_registro = {
        "Número de Registro": contador_registro,
        "Data do Registro": data_registro,
        "Não Conformidade Aberta por": nao_conformidade_aberta_por,
        "Nº Pedido do Cliente": numero_pedido_cliente,
        "Tipo de Não Conformidade": tipo_nao_conformidade,
        "Descreva o Fato": descreva_o_fato,
        "Ação Corretiva Imediata": acao_corretiva_imediata,
        "Responsável pela Ação Corretiva Imediata": responsavel_acao_corretiva
    }
    registros.append(novo_registro)

    # Exibir mensagem de sucesso
    st.success("Registro de Não Conformidade realizado com sucesso!")

    # Carregar o modelo de documento (template)
    doc = Document(template_path)

    # Preencher campos do modelo de documento
    docx_replace(doc, "[CONTADOR_REGISTRO]", contador_registro_text)
    docx_replace(doc, "[DATA_REGISTRO]", data_registro)
    docx_replace(doc, "[NAO_CONFORMIDADE_ABERTA_POR]", nao_conformidade_aberta_por)
    docx_replace(doc, "[NUMERO_PEDIDO_CLIENTE]", numero_pedido_cliente)
    docx_replace(doc, "[TIPO_NAO_CONFORMIDADE]", tipo_nao_conformidade)
    docx_replace(doc, "[DESCREVA_O_FATO]", descreva_o_fato)
    docx_replace(doc, "[ACAO_CORRETIVA_IMEDIATA]", acao_corretiva_imediata)
    docx_replace(doc, "[RESPONSAVEL_ACAO_CORRETIVA]", responsavel_acao_corretiva)

    # Salvar o documento DOCX com nome específico
    filename = f"registros_nao_conformidades_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    # Exibir link para download do arquivo DOCX
    with open(filename, 'rb') as f:
        base64_encoded_docx = base64.b64encode(f.read()).decode()
        href = f"<a href='data:application/octet-stream;base64,{base64_encoded_docx}' download='{filename}'>Baixar Arquivo DOCX</a>"
        st.markdown(href, unsafe_allow_html=True)

    # Limpar os campos do formulário
    nao_conformidade_aberta_por = ""
    numero_pedido_cliente = ""
    tipo_nao_conformidade = ""
    descreva_o_fato = ""
    acao_corretiva_imediata = ""
    responsavel_acao_corretiva = ""

    # Atualizar DataFrame com novo registro
    if df is None:
        df = pd.DataFrame(registros)
    else:
        df = pd.concat([df, pd.DataFrame(registros)])

    # Salvar DataFrame atualizado no arquivo Excel
    df.to_excel('registros_nao_conformidades.xlsx', index=False)

   # Manipulação dos dados e indicadores
    if df is not None:
        df['Data do Registro'] = pd.to_datetime(df['Data do Registro'], format="%d/%m/%Y %H:%M:%S", errors='coerce')
        df['Dia'] = df['Data do Registro'].dt.day
    
        # Registros por Dia
        registros_por_dia = df.groupby(['Ano', 'Mês', 'Dia']).size().reset_index()
    
        # Exibir os indicadores
        st.write("Registros por Dia:")
        st.dataframe(registros_por_dia)
    
        # Gráfico de Registros por Dia
        st.write("Gráfico de Registros por Dia:")
        registros_por_dia_chart = registros_por_dia.copy()
        registros_por_dia_chart['Data'] = registros_por_dia_chart.apply(lambda row: datetime(row['Ano'], row['Mês'], row['Dia']), axis=1)
        registros_por_dia_chart = registros_por_dia_chart.sort_values('Data')
        st.line_chart(registros_por_dia_chart['Data'], registros_por_dia_chart[0])
    
    # Exibir os indicadores
    st.subheader("Indicadores")
    st.write("Registros por Mês:")
    st.dataframe(registros_por_mes)
    st.write("Registros por Ano:")
    st.dataframe(registros_por_ano)
    
    # Gráfico de Registros por Mês
    st.write("Gráfico de Registros por Mês:")
    st.bar_chart(registros_por_mes)
    
    # Gráfico de Registros por Ano
    st.write("Gráfico de Registros por Ano:")
    st.bar_chart(registros_por_ano)
